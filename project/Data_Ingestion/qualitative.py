"""
Phase 2 — Qualitative pipeline.

PDF/DOCX files → text chunks → sentence-transformer embeddings → ChromaDB.
DuckDB `qualitative_chunks` table stores metadata + raw text.

Key design:
  - LLM never sees raw text directly (but Phase 3 search_context() will
    retrieve it via ChromaDB similarity search)
  - All embeddings computed locally via sentence-transformers (no cloud)
  - ChromaDB uses DuckDB persistence (no server process needed)
"""

from __future__ import annotations

import hashlib
import json
import re
import threading
from pathlib import Path
from typing import Optional

import numpy as np

import chromadb
from chromadb.config import Settings as ChromaSettings

# Local imports
from embeddings import encode_texts


# ── ChromaDB singleton ───────────────────────────────────────────────────────

_chroma_lock = threading.Lock()
_chroma_client: Optional[chromadb.PersistentClient] = None


def _get_chroma_client() -> chromadb.PersistentClient:
    """Return a process-wide ChromaDB PersistentClient (lazy, thread-safe)."""
    global _chroma_client
    if _chroma_client is None:
        with _chroma_lock:
            if _chroma_client is None:
                storage_path = Path(__file__).parent / "chromadb_data"
                storage_path.mkdir(parents=True, exist_ok=True)
                _chroma_client = chromadb.PersistentClient(
                    path=str(storage_path),
                    settings=ChromaSettings(anonymized_telemetry=False),
                )
    return _chroma_client


# ── Regex patterns for numerical claims ─────────────────────────────────────

_CURRENCY_RE = re.compile(
    r"(?:Rs\.?\s*|INR\s*|\$\s*|USD\s*|USD\s*)?"
    r"(\d[\d,]*\.?\d*)\s*"
    r"(lakh|lakhs|crore|crores|million|billion|thousand|mn|bn|cr\.?)\b",
    re.I,
)

_PERIOD_RE = re.compile(
    r"(?:FY\s*\d{2,4}|FY\s*\d{2,4}[\-_]Q[1-4]|FY\s*\d{2,4}[\-_]M\d{1,2}"
    r"|CY\s*\d{4}|CY\d{4}"
    r"|Q[1-4]\s*['\u2019]?\d{2}"
    r"|\d{1,2}[\.\-]\d{2}[\.\-]\d{2,4}"
    r"|\d{1,2}[\-\s](?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\-\s]\d{2,4})",
    re.I,
)

_PERCENTAGE_RE = re.compile(r"(\d+\.?\d*)\s*%")

_BARE_LARGE_RE = re.compile(r"\b(\d{6,})\b")

_MULTIPLIER_RE = re.compile(r"(\d+\.?\d*)\s*x\b")


def _context_window(text: str, start: int, end: int, radius: int = 50) -> str:
    """Extract ±radius chars around [start, end), clamped to string bounds."""
    lo = max(0, start - radius)
    hi = min(len(text), end + radius)
    return text[lo:hi]


# ── detect_numerical_claims ──────────────────────────────────────────────────

def detect_numerical_claims(text: str) -> tuple[bool, list[dict]]:
    """
    Scan text for five families of numerical claims.

    Returns (contains_numerical_claim, claims_list) where each claim dict is:
        {number: str, type: str, context: str}
    """
    claims: list[dict] = []

    for pattern, kind in [
        (_CURRENCY_RE, "currency"),
        (_PERIOD_RE, "period"),
        (_PERCENTAGE_RE, "percentage"),
        (_BARE_LARGE_RE, "bare_large_number"),
        (_MULTIPLIER_RE, "multiplier"),
    ]:
        for m in pattern.finditer(text):
            ctx = _context_window(text, m.start(), m.end())
            claims.append({"number": m.group(), "type": kind, "context": ctx})

    has_claim = any(c["type"] != "bare_large_number" for c in claims)
    return has_claim, claims


# ── chunk_text ───────────────────────────────────────────────────────────────

_WORD_RE = re.compile(r"\s+")


def _word_count(text: str) -> int:
    return len(_WORD_RE.split(text.strip()))


def _sentence_boundaries(text: str) -> list[tuple[int, int]]:
    """
    Return list of (start, end) character offsets for each sentence.
    Splits on . ! ? followed by space or end-of-string.
    """
    pattern = re.compile(r"[^.!?]*[.!?]+\s*")
    boundaries: list[tuple[int, int]] = []
    for m in pattern.finditer(text):
        boundaries.append((m.start(), m.end()))
    return boundaries


def _merge_short_chunks(chunks: list[dict], min_words: int = 150) -> list[dict]:
    """
    Merge adjacent chunks that fall below min_words.
    Merges happen in order; merged result appended back to the list.
    """
    if not chunks:
        return chunks
    result: list[dict] = []
    buffer: dict = chunks[0]
    for chunk in chunks[1:]:
        if _word_count(buffer["text"]) < min_words:
            # Accumulate
            overlap = chunk.get("overlap_text", "")
            merged_text = buffer["text"]
            if overlap:
                merged_text += " " + overlap
            # Discard overlap field after merge
            merged_text_clean = merged_text.strip()
            buffer = {
                "text": merged_text_clean,
                "section_path": buffer["section_path"],
            }
        else:
            result.append(buffer)
            buffer = chunk
    result.append(buffer)
    return result


def chunk_text(text: str, section_path: str = "") -> list[dict]:
    """
    Split text into chunks of 150-300 words with the following rules:
      - Never split mid-paragraph
      - Section headings travel with first body paragraph
      - One-sentence overlap between adjacent chunks
      - Short sections (<150 words) merged with next
      - Headings are never emitted as standalone chunks without body text
    """
    # Step 1: Split on blank lines → atomic paragraph units
    raw_paragraphs = re.split(r"\n\s*\n", text)

    # Step 2: Classify as heading or body, group into sections
    sections: list[tuple[str, list[str]]] = []  # (section_path, body_paras)
    current_heading = section_path or ""
    current_bodies: list[str] = []

    for para in raw_paragraphs:
        para = para.strip()
        if not para:
            continue
        stripped = para.strip()
        # Heuristic: short (≤3 words) all-caps or title-case line = heading
        is_heading = (
            len(stripped.split()) <= 5
            and (
                stripped.isupper()
                or (
                    stripped.istitle()
                    and len(stripped) < 80
                    and not stripped[-1] in ".!?"
                )
            )
        )
        if is_heading and len(stripped.split()) <= 5:
            # New heading — save previous section
            if current_bodies or current_heading:
                sections.append((current_heading, current_bodies))
            current_heading = stripped
            current_bodies = []
        else:
            current_bodies.append(para)

    # Flush last section
    if current_bodies or current_heading:
        sections.append((current_heading, current_bodies))

    if not sections:
        return [{"text": text.strip(), "section_path": section_path}]

    chunks: list[dict] = []

    for sec_path, bodies in sections:
        sec_text = "\n\n".join(bodies)
        wc = _word_count(sec_text)

        if wc == 0:
            continue

        if wc <= 300:
            # Small enough for one chunk — save for potential merge
            chunks.append({
                "text": sec_text,
                "section_path": sec_path or section_path,
            })
        else:
            # Split at sentence boundaries
            sentences: list[str] = []
            for m in re.finditer(r"[^.!?]*[.!?]+\s*", sec_text):
                s = m.group().strip()
                if s:
                    sentences.append(s)

            current: list[str] = []
            current_words = 0
            overlap_sentence = ""

            for sent in sentences:
                sent_words = _word_count(sent)
                if current_words + sent_words > 300 and current:
                    # Emit chunk with overlap
                    chunk_text_body = " ".join(current)
                    if overlap_sentence:
                        chunk_text_body += " " + overlap_sentence
                    chunks.append({
                        "text": chunk_text_body.strip(),
                        "section_path": sec_path or section_path,
                    })
                    current = [overlap_sentence]
                    current_words = _word_count(overlap_sentence)
                    overlap_sentence = ""

                current.append(sent)
                current_words += sent_words

                # Carry last sentence as overlap into next chunk
                overlap_sentence = sent

            # Final chunk
            if current:
                chunk_text_body = " ".join(current)
                if overlap_sentence:
                    chunk_text_body += " " + overlap_sentence
                chunks.append({
                    "text": chunk_text_body.strip(),
                    "section_path": sec_path or section_path,
                })

    # Merge short sections
    chunks = _merge_short_chunks(chunks)

    # Assign stable chunk indices
    for i, chunk in enumerate(chunks):
        chunk["chunk_index"] = i

    return chunks


# ── PDF extraction ────────────────────────────────────────────────────────────

def _separate_text_and_tables(blocks: list) -> tuple[list, list]:
    """
    Separate PyMuPDF blocks into text blocks and table blocks.
    Table blocks are those whose bbox width > block bbox height * 2
    (heuristic: wide blocks are table-like).
    """
    text_blocks: list = []
    table_blocks: list = []
    for block in blocks:
        x0, y0, x1, y1 = block["bbox"]
        width = x1 - x0
        height = y1 - y0
        if width > height * 2 and height > 5:
            table_blocks.append(block)
        else:
            text_blocks.append(block)
    return text_blocks, table_blocks


def _classify_text_region(block_text: str, y0: float, page_height: float,
                         block_count_in_page: int) -> tuple[str, str]:
    """
    Classify a text block's region and chunk type.
    Returns (region_type, chunk_type).
    """
    word_count = len(block_text.split())

    # Footnote: near bottom of page, short
    if y0 > page_height * 0.75 and word_count < 30:
        return "inline_annotation", "footnote"

    # Page footer: confidentiality notices (short, near very bottom)
    if y0 > page_height * 0.90 and word_count < 20:
        return "footer_notes", "caption"

    # Table caption: very short block adjacent to a table block
    if word_count < 15 and y0 < page_height * 0.15:
        return "table", "caption"

    # Body narrative
    return "narrative", "narrative"


def extract_pdf_chunks(pdf_path: Path) -> list[dict]:
    """
    Extract all chunks from a PDF file using PyMuPDF (fitz).
    Returns list of chunk dicts (text, region_type, chunk_type, section_path,
    contains_numerical_claim, numerical_claims).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(str(pdf_path))
    all_chunks: list[dict] = []
    global_chunk_index = 0

    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("blocks")
        text_blocks, table_blocks = _separate_text_and_tables(blocks)
        page_height = page.rect.height

        for block in text_blocks:
            x0, y0, x1, y1 = block["bbox"]
            block_text = block["text"].strip()
            if not block_text:
                continue

            region_type, chunk_type = _classify_text_region(
                block_text, y0, page_height, len(blocks)
            )

            # Skip footer notes for now — they may be noise
            if region_type == "footer_notes":
                continue

            has_num, claims = detect_numerical_claims(block_text)

            # Section path = page-level heading heuristic
            section_path = f"page_{page_num}"

            all_chunks.append({
                "text": block_text,
                "region_type": region_type,
                "chunk_type": chunk_type,
                "section_path": section_path,
                "contains_numerical_claim": has_num,
                "numerical_claims": claims,
                "chunk_index": global_chunk_index,
            })
            global_chunk_index += 1

        # Table blocks → one chunk per table
        for block in table_blocks:
            block_text = block["text"].strip()
            if not block_text:
                continue
            has_num, claims = detect_numerical_claims(block_text)
            all_chunks.append({
                "text": block_text,
                "region_type": "table",
                "chunk_type": "table",
                "section_path": f"page_{page_num}",
                "contains_numerical_claim": has_num,
                "numerical_claims": claims,
                "chunk_index": global_chunk_index,
            })
            global_chunk_index += 1

    doc.close()
    return all_chunks


# ── DOCX extraction ─────────────────────────────────────────────────────────────

def extract_docx_tables(docx_path: Path) -> list[dict]:
    """
    Extract all tables from a DOCX file.
    Returns list of table dicts:
        {headers: [...], rows: [[...], ...], section: str, is_structured: bool}
    is_structured=True if first row contains canonical field or period headers.
    """
    from docx import Document
    from canonical_fields import resolve_alias
    from periods import parse_period

    doc = Document(str(docx_path))
    tables: list[dict] = []

    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows[1:]
        ]

        # Check if structured: first row has canonical field names OR period strings
        first_row_values = set(headers)
        canonical_count = sum(
            1 for h in headers if resolve_alias(h) is not None
        )
        period_count = sum(
            1 for h in headers if parse_period(h) is not None
        )
        is_structured = canonical_count >= 2 or period_count >= 2

        tables.append({
            "headers": headers,
            "rows": rows,
            "section": "",
            "is_structured": is_structured,
        })

    return tables


def extract_docx_chunks(docx_path: Path) -> list[dict]:
    """
    Extract narrative chunks from a DOCX file, preserving section hierarchy.
    Returns list of chunk dicts (same schema as PDF extractor).
    """
    from docx import Document

    doc = Document(str(docx_path))
    all_chunks: list[dict] = []
    global_chunk_index = 0
    current_heading = ""

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        # Heading detection
        if "Heading" in style_name or "Title" in style_name:
            current_heading = text
            continue

        word_count = len(text.split())

        # Footnote: very short (<20 words) or starts with asterisk
        if word_count < 20 or text.startswith("*") or text.startswith("†"):
            region_type = "inline_annotation"
            chunk_type = "footnote"
        else:
            region_type = "narrative"
            chunk_type = "narrative"

        has_num, claims = detect_numerical_claims(text)
        section_path = current_heading if current_heading else "document"

        all_chunks.append({
            "text": text,
            "region_type": region_type,
            "chunk_type": chunk_type,
            "section_path": section_path,
            "contains_numerical_claim": has_num,
            "numerical_claims": claims,
            "chunk_index": global_chunk_index,
        })
        global_chunk_index += 1

    return all_chunks


# ── Structured table routing (bifurcate_tables) ─────────────────────────────

def bifurcate_tables(
    docx_path: Path,
    file_id: int,
    entity_id: int,
    conn,  # duckdb connection
) -> list[dict]:
    """
    Route DOCX tables into two buckets:
      - Structured tables with period columns → write to staging_facts
      - Unstructured tables → return as qualitative chunks

    Returns list of qualitative chunks for unstructured tables only.
    """
    import duckdb as _duckdb

    tables = extract_docx_tables(docx_path)
    unstructured_chunks: list[dict] = []

    for table in tables:
        if table["is_structured"]:
            # Attempt to write structured table to staging_facts
            headers = table["headers"]
            rows = table["rows"]

            # Try to detect period columns vs metric columns
            from periods import parse_period
            from canonical_fields import resolve_alias

            period_cols: list[int] = []
            metric_cols: list[int] = []

            for i, h in enumerate(headers):
                if parse_period(h) is not None:
                    period_cols.append(i)
                elif resolve_alias(h) is not None:
                    metric_cols.append(i)

            # Only write if we have clear period + metric columns
            if period_cols and metric_cols:
                try:
                    for row in rows:
                        period_val = row[period_cols[0]] if period_cols else ""
                        period_spec = parse_period(period_val)
                        if not period_spec:
                            continue
                        for mc in metric_cols:
                            canonical = resolve_alias(headers[mc])
                            if not canonical:
                                continue
                            raw_val_str = row[mc]
                            try:
                                raw_val = float(raw_val_str.replace(",", ""))
                            except (ValueError, TypeError):
                                continue
                            conn.execute("""
                                INSERT INTO staging_facts (
                                    staging_id, file_id, entity_id, canonical_field,
                                    period, value_normalised, currency,
                                    original_unit, conversion_factor, conversion_applied,
                                    raw_value, raw_header, row_context
                                ) VALUES (
                                    nextval('staging_id_seq'), ?, ?, ?, ?,
                                    ?, 'INR', 'INR_absolute', 1.0, FALSE,
                                    ?, ?, ?
                                )
                            """, [
                                file_id, entity_id, canonical, period_spec.canonical,
                                raw_val, raw_val, headers[mc], period_val,
                            ])
                    conn.commit()
                except Exception:
                    # If staging write fails, route to qualitative
                    unstructured_chunks.append({
                        "text": f"Table: {', '.join(headers)}",
                        "region_type": "table",
                        "chunk_type": "table",
                        "section_path": table["section"],
                        "contains_numerical_claim": False,
                        "numerical_claims": [],
                        "chunk_index": 0,
                    })
        else:
            # Non-structured → route to ChromaDB
            table_headers = table["headers"]
            table_rows = table["rows"]
            flat_text = f"Table ({', '.join(table_headers)}): " + " | ".join(
                " | ".join(row) for row in table_rows
            )
            has_num, claims = detect_numerical_claims(flat_text)
            unstructured_chunks.append({
                "text": flat_text,
                "region_type": "table",
                "chunk_type": "table",
                "section_path": table["section"],
                "contains_numerical_claim": has_num,
                "numerical_claims": claims,
                "chunk_index": len(unstructured_chunks),
            })

    return unstructured_chunks


# ── Chunk → fact linking ──────────────────────────────────────────────────────

def link_chunks_to_facts(
    chunks: list[dict],
    conn,  # duckdb connection
    entity_id: int,
) -> list[dict]:
    """
    Populate linked_fact_ids, linked_periods, linked_metrics on each chunk
    by scanning numerical_claims and chunk text for matches against live_facts.
    """
    import json as _json

    # Load all live_facts for this entity
    rows = conn.execute("""
        SELECT fact_id, canonical_field, period
        FROM live_facts
        WHERE entity_id = ?
    """, [entity_id]).fetchall()

    if not rows:
        return chunks

    # Index by canonical field (lowercase) and period
    field_index: dict[str, list[str]] = {}
    period_index: dict[str, list[str]] = {}
    all_field_names: dict[str, str] = {}  # lowercase → canonical

    for fact_id, canonical, period in rows:
        key = canonical.lower()
        all_field_names[key] = canonical
        field_index.setdefault(key, []).append(fact_id)
        period_index.setdefault(period, []).append(fact_id)

    for chunk in chunks:
        linked_fact_ids: set[str] = set()
        linked_periods: set[str] = set()
        linked_metrics: set[str] = set()

        # Scan period mentions in claims
        claims = chunk.get("numerical_claims", [])
        if isinstance(claims, str):
            try:
                claims = _json.loads(claims)
            except Exception:
                claims = []

        for claim in claims:
            if claim.get("type") == "period":
                period_str = claim.get("number", "")
                # Match against known periods
                for known_period, fact_ids in period_index.items():
                    if known_period.replace("-", "").replace("_", "") \
                            in period_str.replace("-", "").replace("_", ""):
                        linked_periods.add(known_period)
                        linked_fact_ids.update(fact_ids)

        # Scan chunk text for canonical field keywords
        text_lower = chunk.get("text", "").lower()
        for field_key, canonical in all_field_names.items():
            if field_key in text_lower:
                linked_metrics.add(canonical)
                linked_fact_ids.update(field_index[field_key])

        chunk["linked_fact_ids"] = list(linked_fact_ids)
        chunk["linked_periods"] = list(linked_periods)
        chunk["linked_metrics"] = list(linked_metrics)

    return chunks


# ── ChromaDB storage ─────────────────────────────────────────────────────────

def store_in_chromadb(
    chunks: list[dict],
    entity_slug: str,
    entity_id: int,
    file_id: int,
) -> list[str]:
    """
    Embed text chunks and store in ChromaDB.
    Returns list of ChromaDB document IDs.
    """
    import json as _json

    if not chunks:
        return []

    coll_name = f"chunks_{entity_slug}"
    client = _get_chroma_client()

    collection = client.get_or_create_collection(
        name=coll_name,
        metadata={
            "entity_id": entity_id,
            "entity_slug": entity_slug,
        },
    )

    texts = [c["text"] for c in chunks]

    if not texts:
        return []

    # Encode all texts in one batch call
    embeddings = encode_texts(texts)

    doc_ids = [f"{entity_slug}_chunk_{i}" for i in range(len(chunks))]

    # Prepare metadata (ChromaDB needs JSON-serialisable values)
    metadatas: list[dict] = []
    for c in chunks:
        claims = c.get("numerical_claims", [])
        if isinstance(claims, str):
            try:
                claims = _json.loads(claims)
            except Exception:
                claims = []
        metadatas.append({
            "entity_id": c.get("entity_id", entity_id),
            "file_id": c.get("file_id", file_id),
            "region_type": c.get("region_type", ""),
            "chunk_type": c.get("chunk_type", ""),
            "section_path": c.get("section_path", "") or "",
            "linked_fact_ids": _json.dumps(c.get("linked_fact_ids", [])),
            "linked_periods": _json.dumps(c.get("linked_periods", [])),
            "linked_metrics": _json.dumps(c.get("linked_metrics", [])),
            "contains_numerical_claim": bool(c.get("contains_numerical_claim", False)),
            "numerical_claims": _json.dumps(claims),
            "chunk_index": c.get("chunk_index", 0),
        })

    collection.upsert(
        ids=doc_ids,
        embeddings=embeddings.tolist(),
        documents=texts,
        metadatas=metadatas,
    )

    return doc_ids


# ── Main ingest per file ─────────────────────────────────────────────────────

def ingest_qualitative_file(
    conn,  # duckdb connection
    file_path: Path,
    file_id: int,
    entity_id: int,
    entity_slug: str,
) -> dict:
    """
    Full ingestion pipeline for one qualitative file (PDF or DOCX).
    Returns summary dict with counts.
    """
    import json as _json
    from schema import update_file_state

    suffix = file_path.suffix.lower()
    update_file_state(conn, file_id, "PROCESSING")

    # Step 1: Extract raw chunks by file type
    if suffix == ".pdf":
        raw_chunks = extract_pdf_chunks(file_path)
    elif suffix == ".docx":
        raw_chunks = extract_docx_chunks(file_path)
        # Also handle tables
        table_chunks = extract_docx_tables(file_path)  # already have table chunks
        unstructured_tables = bifurcate_tables(
            file_path, file_id, entity_id, conn
        )
        # Assign sequential indices after raw_chunks
        start_idx = len(raw_chunks)
        for i, tc in enumerate(unstructured_tables):
            tc["chunk_index"] = start_idx + i
        raw_chunks.extend(unstructured_tables)
    else:
        raw_chunks = []

    # Step 2: Apply chunk_text() to each narrative chunk
    processed_chunks: list[dict] = []
    sub_idx = 0
    for chunk in raw_chunks:
        text = chunk["text"]
        section_path = chunk.get("section_path", "")
        if chunk.get("chunk_type") in ("narrative", "footnote"):
            sub_chunks = chunk_text(text, section_path)
            for sc in sub_chunks:
                sc["region_type"] = chunk["region_type"]
                sc["chunk_type"] = chunk["chunk_type"]
                sc["section_path"] = chunk.get("section_path", "")
                sc["contains_numerical_claim"] = chunk.get("contains_numerical_claim", False)
                sc["numerical_claims"] = chunk.get("numerical_claims", [])
                sc["file_id"] = file_id
                sc["entity_id"] = entity_id
                sc["chunk_index"] = sub_idx
                sub_idx += 1
                processed_chunks.append(sc)
        else:
            chunk["file_id"] = file_id
            chunk["entity_id"] = entity_id
            chunk["chunk_index"] = sub_idx
            sub_idx += 1
            processed_chunks.append(chunk)

    # Step 3: detect_numerical_claims() (already done during extraction,
    # but re-run for sub-chunks created by chunk_text)
    for chunk in processed_chunks:
        if chunk.get("chunk_type") in ("narrative", "footnote"):
            has_num, claims = detect_numerical_claims(chunk["text"])
            chunk["contains_numerical_claim"] = has_num
            chunk["numerical_claims"] = claims

    # Step 4: Link chunks to live_facts
    processed_chunks = link_chunks_to_facts(processed_chunks, conn, entity_id)

    # Step 5: Store in ChromaDB
    chroma_doc_ids = store_in_chromadb(
        processed_chunks, entity_slug, entity_id, file_id
    )

    # Step 6: Write metadata to DuckDB qualitative_chunks table
    for i, chunk in enumerate(processed_chunks):
        chroma_id = chroma_doc_ids[i] if i < len(chroma_doc_ids) else None
        conn.execute("""
            INSERT INTO qualitative_chunks (
                chunk_id, file_id, entity_id, chunk_index, region_type, chunk_type,
                section_path, linked_fact_ids, linked_periods, linked_metrics,
                contains_numerical_claim, numerical_claims, raw_text, chroma_document_id
            ) VALUES (
                nextval('staging_id_seq'), ?, ?, ?, ?, ?,
                ?, ?, ?, ?,
                ?, ?, ?, ?
            )
        """, [
            chunk["file_id"], chunk["entity_id"], chunk["chunk_index"],
            chunk["region_type"], chunk["chunk_type"],
            chunk.get("section_path") or None,
            _json.dumps(chunk.get("linked_fact_ids", [])),
            _json.dumps(chunk.get("linked_periods", [])),
            _json.dumps(chunk.get("linked_metrics", [])),
            chunk.get("contains_numerical_claim", False),
            _json.dumps(chunk.get("numerical_claims", [])),
            chunk["text"],
            chroma_id,
        ])
    conn.commit()

    # Step 7: Update file state
    update_file_state(conn, file_id, "LIVE")

    return {
        "file_id": file_id,
        "chunks_extracted": len(raw_chunks),
        "chunks_stored": len(processed_chunks),
        "chroma_doc_ids": len(chroma_doc_ids),
    }


# ── Sibling-file entity resolution ──────────────────────────────────────────

def _resolve_entity_from_siblings(
    pdf_path: Path,
    conn,  # duckdb connection
) -> tuple[int, str]:
    """
    Heuristic: find a sibling CSV/XLSX in the same directory that shares
    the entity name derived from the PDF filename.
    Falls back to creating a new entity.
    """
    import re

    # Derive entity name from PDF filename
    words = re.split(r"[_\-\s]+", pdf_path.stem)
    skip_words = {
        "cim", "deck", "model", "final", "v1", "v2", "v3", "revised",
        "draft", "copy", "export", "overview", "management", "document",
        "confidential", "prospectus", "presentation", "memo",
    }
    kept = [
        w.title() for w in words
        if w.lower() not in skip_words and not re.match(r"^\d{2,4}$", w)
    ]
    candidate_name = " ".join(kept) if kept else pdf_path.stem.replace("_", " ").title()

    # Look for existing entity with matching or similar name
    rows = conn.execute(
        "SELECT entity_id, entity_name, entity_slug FROM entities"
    ).fetchall()

    for eid, ename, eslug in rows:
        if ename.lower().replace(" ", "") == candidate_name.lower().replace(" ", ""):
            return eid, eslug
        if eslug.replace("_", "") in candidate_name.lower().replace(" ", ""):
            return eid, eslug

    # No match → create new entity
    from schema import get_or_create_entity, _slugify
    slug = _slugify(candidate_name)
    entity_id = get_or_create_entity(conn, candidate_name)
    return entity_id, slug


# ── Main pipeline entry point ────────────────────────────────────────────────

def run_qualitative_pipeline(conn) -> dict:
    """
    Scan input_files/ and example_input_files/ for PDF/DOCX files.
    For each file: register, ingest, store embeddings.
    Returns summary dict: {files_processed, chunks_total}.
    """
    import re
    from schema import register_qualitative_file, _slugify

    input_dirs = [
        Path(__file__).parent.parent / "input_files",
        Path(__file__).parent.parent / "example_input_files",
    ]

    qualitative_extensions = {".pdf", ".docx"}
    files: list[Path] = []
    for d in input_dirs:
        if d.exists():
            files.extend(
                p for p in d.iterdir()
                if p.is_file() and p.suffix.lower() in qualitative_extensions
            )

    if not files:
        print("  No PDF or DOCX files found in input_files/ or example_input_files/.")
        return {"files_processed": 0, "chunks_total": 0}

    # Deduplicate by path
    seen: set[Path] = set()
    unique_files: list[Path] = []
    for f in files:
        if f not in seen:
            seen.add(f)
            unique_files.append(f)

    files = unique_files

    print(f"  Found {len(files)} qualitative file(s): "
          + ", ".join(f.name for f in files))

    files_processed = 0
    chunks_total = 0

    for path in files:
        print(f"\n  Processing: {path.name}")
        try:
            entity_id, entity_slug = _resolve_entity_from_siblings(path, conn)
            file_type = path.suffix.lstrip(".").lower()

            file_id = register_qualitative_file(
                conn, entity_id, path.name, str(path), file_type
            )

            result = ingest_qualitative_file(
                conn, path, file_id, entity_id, entity_slug
            )

            print(f"  ✓ {path.name}: {result['chunks_stored']} chunks stored")
            files_processed += 1
            chunks_total += result["chunks_stored"]

        except Exception as e:
            print(f"  ERROR processing {path.name}: {e}")

    return {
        "files_processed": files_processed,
        "chunks_total": chunks_total,
    }
